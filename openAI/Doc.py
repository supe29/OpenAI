from pathlib import Path
import openai
import streamlit as st
from docx import Document

# Set API key
openai.api_key = st.secrets['pass']

st.header("GPT documentation application")
st.subheader("Upload Microsoft Word Document")
doc_files = st.file_uploader("Upload .docx file", type=[
                             "docx"], accept_multiple_files=True)

article_text = st.text_area("Please enter your code")

#temp = st.slider("Temperature", 0.0, 1.0, 0.5)

number = st.number_input('Insert a para number')
st.write('The current para is ', number)

if st.button("Generate"):
    response = openai.Completion.create(
        engine="text-davinci-003",
        # its querry you can change as you want
        prompt="Plz write summary in short:" + article_text,
        max_tokens=3000,
        temperature=0.5
    )
    res = response["choices"][0]["text"]
    st.info(res)

    if doc_files:
        document = Document(doc_files[0])
    else:
        document = Document()

    # Find the specific paragraph by index
    target_paragraph_index = int(number)
    if target_paragraph_index < len(document.paragraphs):
        target_paragraph = document.paragraphs[target_paragraph_index]

        # Insert the response as a new paragraph before the target paragraph
        new_paragraph = target_paragraph.insert_paragraph_before(res)
    else:
        st.warning("Target paragraph index is out of range.")

    # Save the document
    document.save("modified_document.docx")
